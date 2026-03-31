from datetime import timedelta
from typing import List, Optional

from django.contrib.auth.models import AbstractUser
from django.core.validators import FileExtensionValidator
from django.db import models
from django.utils import timezone
import docx
from langchain_core.messages import AIMessage, HumanMessage
import uuid6
from django.db.models import F
from django.conf import settings
from channels.db import database_sync_to_async

        


class User(AbstractUser):
    state = models.CharField(max_length=100 , null=True, blank=True)
    city = models.CharField(max_length=100 , null=True, blank=True)
    occupation = models.CharField(max_length=100 , null=True, blank=True)
    details = models.TextField(null=True, blank=True , max_length=1000)
    
class ChatSession(models.Model):
    session_id = models.UUIDField(primary_key=True, default=uuid6.uuid7, editable=False)
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    title = models.CharField(max_length=100)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    
    def __str__(self):
        return self.title
    
class ChatMessage(models.Model):
    msg_id = models.UUIDField(primary_key=True, default=uuid6.uuid7, editable=False)
    session = models.ForeignKey(ChatSession, on_delete=models.CASCADE)
    role = models.CharField(max_length=100 , choices=[("user" , "user") , ("assistant" , "assistant")])
    content = models.TextField()
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    
    def __str__(self):
        return self.content

    @classmethod
    @database_sync_to_async
    def get_message_history(cls, session_id: str):
        histroy = []
        messages = cls.objects.filter(session=session_id).order_by("created_at")

        for message in messages:
            histroy.append({
                "id" : str(message.msg_id),
                "role" : message.role,
                "content" : message.content,
                "created_at" : message.created_at.isoformat(),
                "docs_refered" : DocumentRefered.get_document_refered(message.msg_id)
            })
        return histroy
    
    @classmethod
    @database_sync_to_async
    def model_chat_context(cls, session_id: str):
        histroy = []
        messages = cls.objects.filter(session=session_id).order_by("created_at")

        for message in messages:
            if message.role == 'user':histroy.append(HumanMessage(content=message.content))
            else:histroy.append(AIMessage(content=message.content))
        return histroy

    
    
class Document(models.Model):
    doc_id = models.UUIDField(primary_key=True, default=uuid6.uuid7, editable=False)
    title = models.CharField(max_length=100)
    content = models.TextField(null=True , blank=True)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    file = models.FileField(upload_to='documents/',validators=[FileExtensionValidator(['pdf','docx', 'txt'])])
    embedding_status = models.CharField(max_length=100 , choices=[("pending" , "pending") , ("embedding" , "embedding") , ("embedding_done" , "embedding_done") , ("embedding_failed" , "embedding_failed")] , default="pending")
    verified = models.BooleanField(default=False)
    
    def __str__(self):
        return self.title

    def save(self , *args , **kwargs):
        if self.file:
            if self.file.name.endswith('.docx'):
                doc = docx.Document(self.file)
                self.content = "\n".join([para.text for para in doc.paragraphs])
            elif self.file.name.endswith('.pdf'):
                pass
            elif self.file.name.endswith('.txt'):
                self.file.seek(0)
                content = self.file.read()
                if isinstance(content, bytes):
                    self.content = content.decode('utf-8')
                else:
                    self.content = content
        super().save(*args, **kwargs)


class OTP(models.Model):
    otp = models.CharField(max_length=6)
    email = models.EmailField()
    otp_type = models.CharField(max_length=20 , choices=[("forget-password" , "forget-password") , ("register" , "register")])
    created_at = models.DateTimeField(auto_now_add=True)
    verified = models.BooleanField(default=False)
    def is_valid(self):
        return self.created_at + timedelta(minutes=5) > timezone.now()



class DocumentRefered(models.Model):
    message = models.ForeignKey(ChatMessage, on_delete=models.CASCADE)
    document = models.ForeignKey(Document, on_delete=models.CASCADE)
    pages = models.JSONField(default=list)

    @classmethod
    def get_document_refered(cls, msg_id: str):
        # 1. Added 'doc_path' to the values query
        references = cls.objects.filter(message=msg_id).values(
            doc_pk=F('document__doc_id'), 
            doc_id=F('document__doc_id'), 
            title=F('document__title'),
            doc_path=F('document__file'), # Assuming the Field name is 'file'
            pages_list=F('pages') 
        )

        grouped_docs = {}

        for ref in references:
            d_id = ref['doc_id']
            
            if d_id not in grouped_docs:
                # 2. Construct the full URL using settings.MEDIA_URL
                # We use .lstrip('/') to avoid double slashes
                relative_path = ref['doc_path']
                full_url = f"{settings.MEDIA_URL}{relative_path}" if relative_path else None

                grouped_docs[d_id] = {
                    "doc_id": str(d_id),
                    "title": ref['title'],
                    "doc_url": full_url, # New parameter added here
                    "pages": []
                }
        
            if ref['pages_list']:
                # Safe handling: ensure pages_list is iterable
                current_pages = ref['pages_list'] if isinstance(ref['pages_list'], list) else []
                combined_pages = set(grouped_docs[d_id]["pages"] + current_pages)
                grouped_docs[d_id]["pages"] = sorted(list(combined_pages))
                
        return list(grouped_docs.values())